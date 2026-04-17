import os
import re
import tempfile
import mimetypes
import uuid
from typing import List, Optional
from datetime import datetime
from zoneinfo import ZoneInfo
import docx
import PyPDF2
import spacy
import phonenumbers
import boto3
from botocore.exceptions import NoCredentialsError
from fastapi import UploadFile
from PyPDF2 import PdfReader, PdfWriter
from dateutil import parser
from dateutil import parser as date_parser

nlp = spacy.load("en_core_web_sm")

s3_client = boto3.client(
    "s3",
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    region_name=os.getenv("AWS_REGION"),
)

EMAIL_REGEX       = r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,10}\b'
BASIC_PHONE_REGEX = r'(?:\+?\d{1,3}[\s\-]?)?(?:\(?\d{2,4}\)?[\s\-]?)?\d{3,5}[\s\-]?\d{3,5}'
LINKEDIN_REGEX    = r'(https?:\/\/(?:www\.)?linkedin\.com\/[A-Za-z0-9\-\_\/]+)'
SKYPE_REGEX       = r'(?i)(?:Skype\s*ID|Skype)\s*[:\-]?\s*([A-Za-z0-9\.\-_]+)'
AADHAAR_REGEX     = r'\b\d{4}\s?\d{4}\s?\d{4}\b'
PAN_REGEX         = r'\b[A-Z]{5}[0-9]{4}[A-Z]\b'
GENDER_REGEX = r'(?i)(?:gender|sex)\s*[:\-]?\s*(Male|Female|Other)\b'
DATE_PATTERNS     = [
    r'(\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[\s\-]+\d{4})\s*(?:to|\-|\–)\s*(Present|\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)[\s\-]+\d{4})',
    r'(\d{1,2}\/\d{4}|\d{4})\s*(?:to|\-|\–)\s*(Present|\d{1,2}\/\d{4}|\d{4})'
]

SKILLS = [
    # --- LANGUAGES (Additions) ---
    "Python", "JavaScript", "TypeScript", "Java", "C", "C++", "C#", "Go", "Rust", "Ruby",
    "PHP", "Perl", "Swift", "Kotlin", "R", "Scala", "Objective-C", "MATLAB", "Shell Scripting",
    "Dart", "Solidity", "GraphQL", "SQL", "NoSQL", "Bash",

    # --- FRONTEND FRAMEWORKS & LIBS ---
    "HTML", "CSS", "SASS", "LESS", "Bootstrap", "Tailwind CSS", "React", "Angular", "Vue.js",
    "Next.js", "Nuxt.js", "Svelte", "jQuery", "Gatsby", "GSAP", "Material UI", "ShadCN", 
    "Chakra UI", "Ant Design", "Redux", "Zustand", "Recoil", "MobX", "Storybook", "Vite",

    # --- BACKEND & FRAMEWORKS ---
    "Node.js", "Express.js", "Django", "Flask", "FastAPI", "Spring Boot", "ASP.NET", 
    "Laravel", "Symfony", "CodeIgniter", "NestJS", "Strapi", "Socket.io", "Hapi.js",

    # --- DATABASES ---
    "PostgreSQL", "MySQL", "MariaDB", "SQLite", "MongoDB", "Cassandra", "Redis", "Elasticsearch",
    "Firebase", "DynamoDB", "CouchDB", "Neo4j", "Snowflake", "BigQuery", "Oracle Database",
    "Microsoft SQL Server", "Supabase", "Prisma", "Sequelize", "Mongoose",

    # --- CLOUD & DEVOPS ---
    "AWS", "Azure", "Google Cloud", "IBM Cloud", "Heroku", "DigitalOcean", "Vercel", "Netlify",
    "Docker", "Kubernetes", "Terraform", "Ansible", "Jenkins", "GitLab CI/CD", "GitHub Actions",
    "CircleCI", "Nginx", "Apache", "Prometheus", "Grafana", "Linux Administration",

    # --- TESTING ---
    "Jest", "Cypress", "Selenium", "Mocha", "Chai", "Puppeteer", "Playwright", "React Testing Library",
    "JUnit", "Pytest", "Postman",

    # --- MOBILE & WEB3 ---
    "React Native", "Flutter", "Android Development", "iOS Development", "Web3.js", "Ethers.js",

    # --- DATA SCIENCE & AI ---
    "Machine Learning", "Deep Learning", "Data Science", "Artificial Intelligence", 
    "Natural Language Processing", "Computer Vision", "Data Analysis", "Data Visualization", 
    "Pandas", "NumPy", "Matplotlib", "Seaborn", "TensorFlow", "Keras", "PyTorch", "Scikit-learn",
    "LangChain", "OpenAI API", "HuggingFace", "Spark", "Hadoop",

    # --- ARCHITECTURE & CONCEPTS ---
    "DSA", "OOP", "REST APIs", "SDLC", "Microservices", "Serverless", "JWT", "OAuth",
    "Agile", "Scrum", "Kanban", "Waterfall", "JIRA", "Confluence", "Git", "GitHub", "GitLab",

    # --- DESIGN & TOOLS ---
    "UI/UX Design", "Figma", "Adobe XD", "Sketch", "Canva", "Photoshop",

    # --- SOFT SKILLS & BUSINESS ---
    "SEO", "SEM", "Google Analytics", "Accounting", "Financial Analysis", "Technical Writing", 
    "Research", "Customer Service", "Sales", "Negotiation", "Public Speaking", "Problem Solving","MS Excel", 
    "MS-PowerPoint", "Team Leadership", "Communication Skills",
]

SECTION_HEADERS = {
    "experience": [
        "experience", "work experience", "professional experience",
        "employment history", "career history", "work history",
        "internship / experience", "internship/experience",
        "internship & experience", "internship", "work & experience",
    ],
    "education": [
        "education", "academic background", "academic qualifications",
        "educational background", "education and training",
    ],
    "projects":       ["projects", "project work", "key projects", "personal projects"],
    "certifications": ["certifications", "certificates", "licenses"],
}

EXTRA_STOP_HEADERS = [
    "internship / experience", "internship/experience", "internship",
    "achievements", "personal details", "declaration",
    "extra curricular", "volunteer", "hobbies", "references",
    "technical skills", "professional summary", "summary",
]


def upload_file_to_s3(file: UploadFile, folder="resumes") -> str:
    try:
        ext = file.filename.split(".")[-1].lower()
        key = f"{folder}/{uuid.uuid4()}.{ext}"
        content_type, _ = mimetypes.guess_type(file.filename)
        if not content_type:
            content_type = "application/octet-stream"

        if ext == "pdf":
            temp_file_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
            file.file.seek(0)
            reader = PdfReader(file.file)
            writer = PdfWriter()
            writer.append_pages_from_reader(reader)
            writer.add_metadata({"/Title": file.filename, "/Author": ""})
            with open(temp_file_path, "wb") as f:
                writer.write(f)
            upload_path = temp_file_path
        else:
            temp_file_path = tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}").name
            file.file.seek(0)
            with open(temp_file_path, "wb") as f:
                f.write(file.file.read())
            upload_path = temp_file_path

        with open(upload_path, "rb") as f:
            s3_client.upload_fileobj(
                f,
                os.getenv("AWS_BUCKET_NAME"),
                key,
                ExtraArgs={
                    "ContentType": content_type,
                    "ContentDisposition": f'inline; filename="{file.filename}"',
                },
            )
        os.remove(temp_file_path)
        url = f"https://{os.getenv('AWS_BUCKET_NAME')}.s3.{os.getenv('AWS_REGION')}.amazonaws.com/{key}"
        return url

    except NoCredentialsError:
        print("AWS credentials not available")
        return None
    except Exception as e:
        print(f"Failed to upload to S3: {e}")
        return None


def extract_text(file: UploadFile) -> str:
    try:
        if file.filename.lower().endswith(".pdf"):
            file.file.seek(0)
            try:
                reader = PyPDF2.PdfReader(file.file, strict=False)
            except Exception as e:
                print(f"[ERROR] Cannot read PDF: {file.filename} -> {e}")
                return ""
            full_text = []
            for page in reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text:
                        lines = [line.strip() for line in page_text.split("\n") if line.strip()]
                        full_text.extend(lines)
                except Exception as e:
                    print(f"[ERROR] Failed reading PDF page in {file.filename}: {e}")
            return "\n".join(full_text)

        elif file.filename.lower().endswith(".docx"):
            file.file.seek(0)
            try:
                doc = docx.Document(file.file)
            except Exception as e:
                print(f"[ERROR] Failed to parse DOCX file: {file.filename}, error: {e}")
                return ""
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            return "\n".join(full_text)

        elif file.filename.lower().endswith(".txt"):
            file.file.seek(0)
            return file.file.read().decode("utf-8", errors="ignore")

    except Exception as e:
        print(f"[ERROR] extract_text crashed for {file.filename}: {e}")
        return ""

    return ""


def sanitize_text(text: str) -> str:
    if not text:
        return ""
    return text.replace('\x00', '').strip()


def extract_name(text: str) -> Optional[str]:
    if not text:
        return None
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    blacklist = {'resume', 'cv', 'curriculum', 'developer', 'engineer', 'consultant', 'manager'}
    for line in lines:
        match = re.search(r"(?i)^name\s*[:\-]\s*([A-Za-z\s.]+)$", line)
        if match:
            candidate = match.group(1).strip()
            if 2 <= len(candidate.split()) <= 4:
                return candidate
    if lines:
        first_line = lines[0]
        words = first_line.split()
        if (2 <= len(words) <= 4 and
                all(w[0].isupper() for w in words if w.isalpha()) and
                all(w.lower() not in blacklist for w in words) and
                not any(char.isdigit() for char in first_line)):
            return first_line
    for line in lines[:5]:
        words = line.split()
        if (2 <= len(words) <= 4 and
                all(w[0].isupper() for w in words if w.isalpha()) and
                all(w.lower() not in blacklist for w in words) and
                not any(char.isdigit() for char in line)):
            return line
    doc = nlp(text[:800])
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text.strip()
            if (2 <= len(name.split()) <= 4 and
                    not any(char.isdigit() for char in name) and
                    not any(word.lower() in blacklist for word in name.split())):
                return name
    return None


def extract_phone_numbers(text: str) -> List[str]:
    numbers = []
    for match in phonenumbers.PhoneNumberMatcher(text, None):
        try:
            formatted = phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.E164)
            if formatted not in numbers:
                numbers.append(formatted)
        except Exception:
            continue
    return numbers

def extract_section(text: str, section_names: List[str]) -> List[str]:
    lines = text.splitlines()
    cleaned = [l.strip() for l in lines if l.strip()]
    joined_text = "\n".join(cleaned)

    pattern = '|'.join([re.escape(h) for h in section_names])
    split = re.split(rf"(?im)^[ \t]*({pattern})[ \t]*$", joined_text)

    if len(split) < 3:
        split = re.split(rf"(?i)({pattern})", joined_text)
        if len(split) < 3:
            return []

    content = split[2]

    all_stop_headers = (
        [h for headers in SECTION_HEADERS.values() for h in headers]
        + EXTRA_STOP_HEADERS
    )
    stop_pattern = '|'.join([re.escape(h) for h in all_stop_headers])
    stop_match = re.search(rf"(?i)({stop_pattern})", content)
    if stop_match:
        content = content[:stop_match.start()]

    return [l.strip()
     for l in content.strip().splitlines() if l.strip()]


def extract_skills(text: str) -> List[str]:
    lower_text = text.lower()
    found = set()
    for skill in SKILLS:
        if re.search(rf'\b{re.escape(skill.lower())}\b', lower_text):
            found.add(skill)
    return list(found)


def extract_education_section(text: str) -> List[str]:
    raw_lines = extract_section(text, SECTION_HEADERS["education"])
    cleaned = []
    for line in raw_lines:
        line = re.sub(r'^[•\-\*\u2022]+\s*', '', line)
        line = re.sub(r'Education[:\-\s]*', '', line, flags=re.I)
        if line.strip():
            cleaned.append(line.strip())
    return cleaned


def normalize_date(date_str):
    if not date_str:
        return None
    date_str = date_str.strip().replace('.', '')
    if date_str.lower() in ("present", "current"):
        return None
    try:
        dt = date_parser.parse(date_str)
        return f"{dt.year}-{dt.month:02d}"
    except Exception:
        return None


def extract_date_range(text):
    for pattern in DATE_PATTERNS:
        m = re.search(pattern, text, re.IGNORECASE)
        if m:
            start, end = m.groups()
            return normalize_date(start), normalize_date(end)
    return None, None


def parse_resume_text(text: str):
    name = extract_name(text)
    emails = re.findall(EMAIL_REGEX, text)
    email = emails[0] if emails else None
    phones = extract_phone_numbers(text)
    if not phones:
        fallback_phones = [
            p for p in re.findall(BASIC_PHONE_REGEX, text)
            if len(re.sub(r'\D', '', p)) >= 10
        ]
        phones = fallback_phones
    phone = phones[0] if phones else None
    skills = extract_skills(text)
    experience = extract_section(text, SECTION_HEADERS["experience"])
    all_lines = [l.strip() for l in text.splitlines() if l.strip()]
    work_experiences = parse_work_experiences(all_lines)
    education = extract_education_section(text)
    projects = extract_section(text, SECTION_HEADERS["projects"])
    certifications = extract_section(text, SECTION_HEADERS["certifications"])
    linkedin_match = re.search(LINKEDIN_REGEX, text)
    skype_match = re.search(SKYPE_REGEX, text)
    aadhaar_match = re.search(AADHAAR_REGEX, text)
    pan_match = re.search(PAN_REGEX, text)
    gender_match = re.search(GENDER_REGEX, text)

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "skills": skills,
        "experience": experience,
        "work_experiences": work_experiences,
        "education": education,
        "projects": projects,
        "certifications": certifications,
        "linkedin": linkedin_match.group(0) if linkedin_match else None,
        "skypeId": skype_match.group(1) if skype_match else None,
        "aadharCardNumber": aadhaar_match.group(0) if aadhaar_match else None,
        "panCardNumber": pan_match.group(0) if pan_match else None,
        "gender": gender_match.group(1).capitalize() if gender_match else None,
        "raw_text": sanitize_text(text),
    }


def split_name(full_name: Optional[str]):
    if not full_name:
        return None, None
    parts = full_name.strip().split()
    if len(parts) == 1:
        return parts[0], ""
    return parts[0], " ".join(parts[1:])



def parse_education_items(education_lines: List[str]):
    edu_list = []
    i = 0
    
    degree_map = {
        r"\bB\.?TECH\b": "Bachelor of Technology",
        r"\bB\.?E\.?\b": "Bachelor of Engineering",
        r"\bM\.?TECH\b": "Master of Technology",
        r"\bM\.?E\.?\b": "Master of Engineering",
        r"\bMCA\b": "Master of Computer Applications",
        r"\bBCA\b": "Bachelor of Computer Applications",
        r"\bB\.?SC\s+I\.?T\b": "B.Sc. in Information Technology",
        r"\bM\.?SC\s+I\.?T\b": "M.Sc. in Information Technology",
        r"\bMBA\b": "Master of Business Administration",
        r"\bB\.?COM\b": "Bachelor of Commerce",
        r"\bB\.?SC\b": "Bachelor of Science",
        r"\bM\.?SC\b": "Master of Science",
        r"\bC\.?A\.?\b": "Chartered Accountant", 
        r"\bPH\.?D\b": "Doctor of Philosophy",
        r"\bB\.?A\.?\b": "Bachelor of Arts",
        r"\bM\.?A\.?\b": "Master of Arts",
        r"\bINTERMEDIATE\b|\b12TH\b|\bHSC\b": "Intermediate (12th)",
        r"\b10TH\b|\bHIGHSCHOOL\b|\bSSC\b": "High School (10th)",
    }

    while i < len(education_lines):
        line = education_lines[i].strip()
        
        if not line or any(kw in line.lower() for kw in ["internship", "experience", "---"]):
            i += 1
            continue

        context_window = " ".join(education_lines[i:i+3]).replace('\n', ' ')
        
        degree = ""
        for pattern, full_name in degree_map.items():
            if re.search(pattern, context_window, re.I):
                degree = full_name
                break

        uni_match = re.search(r'([^,0-9]+(?:University|Institute|College|Academy|School))', context_window, re.I)
        uni_name = uni_match.group(1).strip() if uni_match else ""

        if degree or uni_name:
            edu_type = "GRADUATION"
            degree_upper = degree.upper()
            
            if any(x in degree_upper for x in ["MASTER", "M.TECH", "MBA", "MCA", "M.SC", "M.A", "M.E"]):
                edu_type = "POST_GRADUATION"
            elif "PH.D" in degree_upper or "PHILOSOPHY" in degree_upper:
                edu_type = "DOCTORATE"
            elif any(x in degree_upper for x in ["10TH", "12TH", "SCHOOL", "INTERMEDIATE", "SSC", "HSC"]):
                edu_type = "SCHOOLING"
            elif not degree and "School" in uni_name:
                edu_type = "SCHOOLING"

            years = re.findall(r'\b(20\d{2}|19\d{2})\b', context_window)
            year = max(years) if years else ""

            edu_list.append({
                "educationId": None,
                "type": edu_type,
                "institutionName": uni_name,
                "passingYear": year,
                "degreeName": degree,
                "board": "",
                "percentage": "",
                "university": uni_name,
                "createdAt": datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S"),
                "updatedAt": datetime.now(ZoneInfo("Asia/Kolkata")).strftime("%Y-%m-%d %H:%M:%S"),
                "educationStatus": "COMPLETED",
            })
            i += 2 
        else:
            i += 1

    return edu_list


def parse_work_experiences(lines, candidateId=None):
    work_experiences = []

    date_range_pattern = re.compile(
        r'(?P<start>(?:January|February|March|April|May|June|July|August|September|October|November|December|'
        r'Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s?\d{4})\s*[-–\u2014to/]+\s*'
        r'(?P<end>Present|Current|'
        r'(?:January|February|March|April|May|June|July|August|September|October|November|December|'
        r'Jan|Feb|Mar|Apr|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s?\d{4})',
        re.I,
    )

    action_verbs = [
        "Managed", "Led", "Developed", "Prepared", "Assisted",
        "Monitored", "Handled", "Maintained", "Conducted"
    ]

    experience_headers = [
        "work experience", "professional experience", "employment history",
        "career history", "work history", "experience",
        "internship experience", "internship & experience",
        "internship/experience", "internship / experience",
        "internship",
    ]

    stop_headers = [
        "skills", "education", "certifications", "projects",
        "achievements", "personal details", "declaration",
        "extra curricular", "volunteer", "hobbies", "references",
        "technical skills", "professional summary", "summary",
    ]

    def _normalize(d):
        if not d or str(d).lower() in ("present", "current"):
            return None
        try:
            dt = date_parser.parse(d)
            return f"{dt.year}-{dt.month:02d}"
        except:
            return None

    def _is_bullet_or_description(line):
        return (
            line.startswith(("•", "-", "*", "·")) or
            any(line.lower().startswith(v.lower()) for v in action_verbs)
        )

    def _is_experience_header(line):
        return any(line.lower().strip() == h for h in experience_headers)

    def _is_stop_header(line):
        return any(line.lower().strip() == h for h in stop_headers)

    in_experience = False
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        if _is_experience_header(line):
            in_experience = True
            i += 1
            continue

        if in_experience and _is_stop_header(line):
            in_experience = False
            i += 1
            continue

        if not in_experience:
            i += 1
            continue

        dr = date_range_pattern.search(line)
        if dr:
            start, end = dr.group("start"), dr.group("end")
            remaining_text = date_range_pattern.sub("", line).strip(", |–-\u2014").strip()

            context = []
            back_idx = i - 1
            while back_idx >= 0 and len(context) < 3:
                prev_line = lines[back_idx].strip()

                if not prev_line:
                    back_idx -= 1
                    continue

                if _is_experience_header(prev_line) or _is_stop_header(prev_line):
                    break

                if _is_bullet_or_description(prev_line):
                    break

                context.append(prev_line)
                back_idx -= 1

            role = ""
            company = ""

            if len(context) >= 2:
                company = context[0]
                role = context[1]
            elif len(context) == 1:
                company = context[0]

            if remaining_text:
                if not company:
                    company = remaining_text
                elif remaining_text.lower() not in company.lower():
                    company = f"{company}, {remaining_text}"

            work_experiences.append({
                "candidateId": candidateId,
                "workExperienceId": None,
                "role": role,
                "companyName": company,
                "startDate": _normalize(start),
                "endDate": _normalize(end),
                "isCurrentlyWorking": bool(end and end.lower() in ("present", "current")),
            })

        i += 1

    return work_experiences

    
def parse_projects(lines, candidateId=None):
    projects = []
    proj_id = 1
    current_proj = None

    date_pattern = re.compile(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\.?\s*\d{4}', re.I)
    tools_pattern = re.compile(r'Tools\s+Used\s*:\s*(.+)', re.I)

    bad_lines = {
        "technical skills", "extra curricular", "strength and weakness",
        "technology stacks", "last updated", "page", "programming languages",
        "design and development tools", "web technologies", "databases",
        "cloud platforms", "volunteer work", "reading books", "education",
        "certifications", "interests", "hobbies"
    }

    def _normalize(d):
        if not d or str(d).lower() in ("present", "current"):
            return None
        try:
            dt = date_parser.parse(d)
            return f"{dt.year}-{dt.month:02d}"
        except Exception:
            return None

    def clean_text(line):
        return re.sub(r'\s+', ' ', line).strip()

    def save_current():
        nonlocal current_proj, proj_id
        if not current_proj:
            return
        
        name = current_proj.get("projectName")
        desc_list = current_proj.get("description", [])
        
        if not name or (not desc_list and not current_proj.get("technologies")):
            current_proj = None
            return

        description = " ".join(desc_list).strip()
        
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        projects.append({
            "candidateId": candidateId,
            "projectId": None,
            "projectName": name,
            "client": None,
            "clientLocation": None,
            "startDate": _normalize(current_proj.get("startDate")),
            "endDate": _normalize(current_proj.get("endDate")),
            "role": None,
            "technologies": current_proj.get("technologies", []),
            "description": description if description else None,
            "projectIndustry": None,
            "duration": None,
            "teamSize": None,
            "locationMode": None,
            "projectUrl": None,
            "projectImages": [],
            "createdAt": now,
            "updatedAt": now,
        })
        proj_id += 1
        current_proj = None

    def is_heading(line, raw_line):
        l = line.lower()
        if any(bad in l for bad in bad_lines):
            return False
        
        if len(l.split()) > 12:
            return False
            
        action_verbs = ("developed", "built", "created", "implemented", "used", "designed", "integrated")
        if any(l.startswith(v) for v in action_verbs):
            return False

        if line.strip().endswith(".") and len(l.split()) > 4:
            return False

        return 1 < len(l.split())

    i = 0
    while i < len(lines):
        raw_line = lines[i]
        line = clean_text(raw_line.strip("•- \t"))
        
        if not line:
            i += 1
            continue

        lower = line.lower()

        if any(bad in lower for bad in bad_lines):
            save_current()
            i += 1
            continue

        m = tools_pattern.search(line)
        if m:
            if not current_proj:
                current_proj = {"projectName": f"Project {proj_id}", "startDate": None, "endDate": None, "technologies": [], "description": []}
            current_proj["technologies"] = [t.strip() for t in m.group(1).split(",") if t.strip()]
            i += 1
            continue

        if date_pattern.fullmatch(line):
            if current_proj:
                if not current_proj.get("startDate"):
                    current_proj["startDate"] = line
                else:
                    current_proj["endDate"] = line
            i += 1
            continue

        if is_heading(line, raw_line):
            save_current() 
            current_proj = {"projectName": line, "startDate": None, "endDate": None, "technologies": [], "description": []}
            
            dm = date_pattern.search(line)
            if dm:
                current_proj["startDate"] = dm.group()
            i += 1
            continue

        if current_proj:
            text = "GitHub Link" if "github.com" in lower else line
            current_proj["description"].append(text)
        
        i += 1

    save_current() 
    return projects


def calculate_parsability_score(text: str) -> dict:
    if not text.strip():
        return {"score": 0.0, "details": {}, "parsable": False}

    checks = {
        "name": bool(extract_name(text)),
        "email": bool(re.search(EMAIL_REGEX, text)),
        "phone": bool(extract_phone_numbers(text)),
        "skills": bool(extract_skills(text)),
        "education": bool(extract_education_section(text)),
        "experience": bool(extract_section(text, SECTION_HEADERS["experience"])),
    }

    score = (sum(1 for v in checks.values() if v) / len(checks)) * 100
    parsable = score >= 80

    return {
        "score":    round(score, 2),
        "details":  checks,
        "parsable": parsable,
    }