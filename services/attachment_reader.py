# services/attachment_reader.py
import re
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import openpyxl
import pandas as pd
from docx import Document

SUPPORTED_TYPES = ["pdf", "docx", "xlsx"]


# ── PDF Reader ────────────────────────────────────────────────
def read_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_number, page in enumerate(pdf.pages):
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += f"\n--- Page {page_number + 1} ---\n"
                    text += page_text + "\n"

                tables = page.extract_tables()
                if tables:
                    text += f"\n--- Tables on Page {page_number + 1} ---\n"
                    for table in tables:
                        for row in table:
                            clean_row = [cell if cell else "" for cell in row]
                            text += " | ".join(clean_row) + "\n"

        if not text.strip():
            print(f"No text found in {file_path} — trying OCR...")
            images = convert_from_path(file_path)
            for page_number, image in enumerate(images):
                print(f"Running OCR on page {page_number + 1}...")
                page_text = pytesseract.image_to_string(image)
                if page_text.strip():
                    text += f"\n--- Page {page_number + 1} (OCR) ---\n"
                    text += page_text + "\n"

        return text.strip()

    except FileNotFoundError:
        print(f"PDF file not found: {file_path}")
        return ""
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
        return ""


# ── Word Reader ───────────────────────────────────────────────
def read_word(file_path):
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        for table_number, table in enumerate(doc.tables):
            text += f"\n--- Table {table_number + 1} ---\n"
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text += row_text + "\n"
        return text.strip()
    except FileNotFoundError:
        print(f"Word file not found: {file_path}")
        return ""
    except Exception as e:
        print(f"Error reading Word file {file_path}: {e}")
        return ""


# ── Excel Reader ──────────────────────────────────────────────
def read_excel(file_path):
    text = ""
    try:
        workbook = openpyxl.load_workbook(file_path)
        sheet_names = workbook.sheetnames
        print(f"Found {len(sheet_names)} sheet(s) in {file_path}")
        for sheet_name in sheet_names:
            text += f"\n--- Sheet: {sheet_name} ---\n"
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            df = df.dropna(how="all")
            df = df.dropna(axis=1, how="all")
            text += df.to_string(index=False)
            text += "\n"
        return text.strip()
    except FileNotFoundError:
        print(f"Excel file not found: {file_path}")
        return ""
    except Exception as e:
        print(f"Error reading Excel file {file_path}: {e}")
        return ""


# ── Excel Summary ─────────────────────────────────────────────
def get_excel_summary(file_path):
    summary = {}
    try:
        workbook = openpyxl.load_workbook(file_path)
        summary["total_sheets"]  = len(workbook.sheetnames)
        summary["sheet_names"]   = workbook.sheetnames
        summary["sheets_detail"] = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            summary["sheets_detail"].append({
                "name"         : sheet_name,
                "total_rows"   : sheet.max_row,
                "total_columns": sheet.max_column
            })
        return summary
    except Exception as e:
        print(f"Error getting Excel summary: {e}")
        return {}


# ── Read Any Attachment ───────────────────────────────────────
def read_attachment(file_path: str) -> str:
    ext = file_path.split(".")[-1].lower()
    if ext not in SUPPORTED_TYPES:
        print(f"Skipping unsupported file type: {ext}")
        return ""
    if ext == "pdf":
        return read_pdf(file_path)
    elif ext == "docx":
        return read_word(file_path)
    elif ext == "xlsx":
        return read_excel(file_path)
    return ""


def extract_from_attachment_text(text: str) -> dict:
    result = {
        "phone"     : None,
        "linkedin"  : None,
        "github"    : None,
        "skills"    : [],
        "experience": None,
    }

    if not text:
        return result

    # ── Phone number ──────────────────────────────────────────
    phone_match = re.search(
        r'(\+?\d{1,3}[\s\-]?)?(\(?\d{3}\)?[\s\-]?)(\d{3}[\s\-]?\d{4})',
        text
    )
    if phone_match:
        result["phone"] = phone_match.group(0).strip()

    # ── LinkedIn ──────────────────────────────────────────────
    linkedin_match = re.search(
        r'linkedin\.com/in/([A-Za-z0-9\-_]+)', text, re.IGNORECASE
    )
    if linkedin_match:
        result["linkedin"] = f"https://linkedin.com/in/{linkedin_match.group(1)}"

    # ── GitHub ────────────────────────────────────────────────
    github_match = re.search(
        r'github\.com/([A-Za-z0-9\-_]+)', text, re.IGNORECASE
    )
    if github_match:
        result["github"] = f"https://github.com/{github_match.group(1)}"

    # ── Skills ────────────────────────────────────────────────
    skill_keywords = [
        "python", "java", "javascript", "typescript", "react", "node",
        "angular", "vue", "flutter", "django", "fastapi", "flask",
        "devops", "machine learning", "deep learning", "data science",
        "backend", "frontend", "fullstack", "android", "ios", "php",
        "golang", "rust", "c++", "c#", "dotnet", ".net", "aws", "azure",
        "gcp", "docker", "kubernetes", "sql", "postgresql", "mongodb",
        "redis", "elasticsearch", "kafka", "spark", "hadoop", "tableau",
        "power bi", "excel", "git", "linux", "bash", "tensorflow",
        "pytorch", "opencv", "nlp", "qa", "selenium", "testing"
    ]
    found_skills = []
    for skill in skill_keywords:
        if re.search(rf'\b{re.escape(skill)}\b', text, re.IGNORECASE):
            found_skills.append(skill.title())
    result["skills"] = found_skills

    # ── Years of Experience ───────────────────────────────────
    exp_match = re.search(
        r'(\d+)\+?\s*(?:years?|yrs?)[\s\w]*(?:of\s+)?experience',
        text, re.IGNORECASE
    )
    if exp_match:
        result["experience"] = f"{exp_match.group(1)} years"

    return result



def process_attachment(file_path: str) -> dict:
    text            = read_attachment(file_path)
    extracted       = extract_from_attachment_text(text)
    extracted["raw_text"] = text  
    return extracted